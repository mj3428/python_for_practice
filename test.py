#!/usr/bin/env python
# encoding: utf-8
'''
@author: miaojue
@contact: major3428@foxmail.com
@software: pycharm
@file: test.py
@time: 2018-10-16 下午 2:54
@desc:
'''

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches
import numpy as np
import pandas as pd

#公司名
com_name = '嘉名染整'
document = Document()
for i in range(0,5):
    document.add_paragraph('')
pic = document.add_picture('./pic/poweryun.png', height=Cm(3.43), width=Cm(9.83))
#南德电气1.82×6.07 电能卫士3.43×9.83
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中设置

document.styles['Heading 1'].font.name = u'宋体'
document.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Heading 1'].font.size = Pt(12)

paragraph = document.add_heading(text=com_name+'专项数据体检报告', level=1)
document.save('./text/test.docx')
