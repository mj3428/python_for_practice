#!/usr/bin/env python
# encoding: utf-8
'''
@author: miaojue
@contact: major3428@foxmail.com
@software: pycharm
@file: test.py
@time: 2018-10-16 下午 2:54
@desc:
'''

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import Inches
import numpy as np
import pandas as pd
from docx.oxml.ns import qn
import datetime

#参数配置
com_name = '嘉名染整' #公司名
transformer = '地下室 1#' #变压器名
now = datetime.datetime.now() #当前时间的datetime
document = Document()

#添加样式

#开始写
for i in range(0,5):
    document.add_paragraph('')
pic = document.add_picture('./pic/poweryun.png', height=Cm(3.43), width=Cm(9.83))
#南德电气1.82×6.07 电能卫士3.43×9.83
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中设置

style_T1 = document.styles.add_style('T1', WD_STYLE_TYPE.PARAGRAPH) #创建T1样式
style_T1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
document.styles['T1'].font.name = u'宋体' #T1样式使用字体
document.styles['T1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['T1'].font.size = Pt(26)
document.styles['T1'].font.bold = True

style_T2 = document.styles.add_style('T2', WD_STYLE_TYPE.PARAGRAPH) #创建T2样式
style_T2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
document.styles['T2'].font.name = u'宋体' #T2样式使用字体
document.styles['T2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['T2'].font.size = Pt(22)
document.styles['T2'].font.bold = True

document.add_paragraph(text='\n\n'+com_name+'有限公司', style='T1')
document.add_paragraph(text=transformer+'变压器\n电能健康分析评估报告\n', style='T2')
#插入印章
#picture = document.add_picture('./pic/poweryun_seal.png', height=Cm(4.00), width=Cm(4.32))
#创建日期样式
style_D1 = document.styles.add_style('Date1', WD_STYLE_TYPE.PARAGRAPH) #创建Date1样式
style_D1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
document.styles['Date1'].font.name = u'宋体' #T1样式使用字体
document.styles['Date1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Date1'].font.size = Pt(14)
p = document.add_paragraph(style='Date1')
r = p.add_run() #可以通过run来继续填写内容，包括图片、文字
r.add_picture('./pic/poweryun_seal.png',  height=Cm(4.00), width=Cm(4.32))
r.add_text(now.strftime('%Y')+'年'+now.strftime('%m')+'月')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 图片和文字右对齐设置

#document.add_paragraph(text=now.strftime('%Y')+'年'+now.strftime('%m')+'月', style='Date1')
document.add_page_break() #插入分页符
#标题不能改 可用add_style添加一个后替换
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Normal'].font.size = Pt(10.5)


document.save('./text/test.docx')
