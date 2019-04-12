#!/usr/bin/env python
# encoding: utf-8
'''
@author: miaoj
@contact: major3428@foxmail.com
@software: pycharm
@file: test.py
@time: 2018-10-16 下午 2:54
@desc:
'''

from docx import Document
from docx.shared import Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import numpy as np
import pandas as pd
from docx.oxml.ns import qn
import datetime
from auto_config import *
from auto_cal import Calculate
from auto_cal import Talk
import time


start = time.clock()#计算导入数据时间
calc = Calculate()
talk = Talk()
calc.add_model()
quality = calc.quality()
max_value = calc.calculate()
max_value = np.around(max_value, 2)
conclusion = (talk.utalk(quality[0]), talk.uthdtalk(quality[1]), talk.lftalk(max_value[5]),
              talk.ithdtalk(calc.ithd_risk), talk.pftalk(calc.pf_mean),
              talk.lftalk(max_value[5]), talk.unbtalk(quality[5]))
riskamount = calc.count_risk()
qramount = calc.qramount
max_trend = calc.group()
min_trend = calc.min_trend
res_color = talk.result
print(res_color)
resDic = {'green': '—',
          'yellow': '↑',
          'red': '↑',}
brush = {'green': RGBColor(0x00, 0xB0, 0x50),
         'yellow': RGBColor(0xff, 0xff, 0x00),
         'red': RGBColor(0xff, 0x00, 0x00)}
chapter = ('电压', '谐波电压', '电流', '谐波电流', '功率因数', '负荷率', '三相电流不平衡度')
risk_list = []
result = []
brush_color = []
for i in res_color:
    result.append(resDic[i])
for i in res_color:
    brush_color.append(brush[i])
ele_n = KVA * 1.4434
now = datetime.datetime.now() #当前时间的datetime
document = Document(NANDECLOUD_MODEL)
#至此时间消耗:1.16-1.32s


#开始写入
style_T1 = document.styles.add_style('T1', WD_STYLE_TYPE.PARAGRAPH) #创建T1样式
style_T1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
document.styles['T1'].font.name = u'宋体' #T1样式使用字体
document.styles['T1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['T1'].font.size = Pt(26)
document.styles['T1'].font.bold = True

style_T2 = document.styles.add_style('T2', WD_STYLE_TYPE.PARAGRAPH) #创建T2样式
style_T2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
document.styles['T2'].font.name = u'宋体' #T2样式使用字体
document.styles['T2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['T2'].font.size = Pt(22)
document.styles['T2'].font.bold = True

document.paragraphs[14].style = 'T1'
document.paragraphs[14].add_run(text= COM_NAME+'有限公司')
document.paragraphs[16].clear()
document.paragraphs[16].style = 'T2'
document.paragraphs[16].add_run(text='('+ MONTH +TRANSFORMER+'变压器)')
#document.paragraphs[18].style = 'T2'
#document.paragraphs[18].add_run(text='电能健康分析评估报告\n')

#插入印章
#picture = document.add_picture('./pic/poweryun_seal.png', height=Cm(4.00), width=Cm(4.32))
#创建日期样式
style_D1 = document.styles.add_style('Date1', WD_STYLE_TYPE.PARAGRAPH) #创建Date1样式
style_D1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
document.styles['Date1'].font.name = u'宋体' #T1样式使用字体
document.styles['Date1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.styles['Date1'].font.size = Pt(14)
document.paragraphs[25].clear()
document.paragraphs[25].style = 'Date1'
document.paragraphs[25].add_run(text=now.strftime('%Y')+'年 '+now.strftime('%m')+'月')
#创建新样式消耗0.01S

last = document.paragraphs[-1]
last.style = 'Heading 1'
last.add_run(text='一、监测概况及结论')
document.add_heading('1.1电能参数体检结果', level=2)

records = (('电压数据', result[0], str(max_value[0]) + 'V', '205~235V', '《GB/T 12325-2008》'),
          ('谐波电压数据', result[1], str(max_value[1]) + '%%', '＜5%', '《GB/T14549-1993》'),
          ('电流数据', result[2], str(max_value[2]) + 'A', '＜%.1fA' % ele_n, '《JGJ16-2008》'),
          ('谐波电流数据', result[3], str(max_value[3]) + 'A', '＜各分次国标限值', '《GB/T14549-1993》'),
          ('功率因数', result[4], str(max_value[4]) + '(Min)', '0.9~1.0', '《JGJ16-2008》'),
          ('负荷率', result[5], str(max_value[5]) + '%', '＜额定容量85%', '《JGJ16-2008》'),
          ('三相电流不平衡度', result[6], str(max_value[6]) + '%', '＜15%', '《GB/T 1094-2013》'))
table1 = document.add_table(rows=1, cols=5)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = '体检项目'
hdr_cells[1].text = '体检结果'
hdr_cells[2].text = '体检值(Max)'
hdr_cells[3].text = '参考值'
hdr_cells[4].text = '参考标准'
for sty, res, mv, ran, txt in records:
    row_cells = table1.add_row().cells
    row_cells[0].text = sty
    row_cells[1].text = res
    row_cells[2].text = mv
    row_cells[3].text = ran
    row_cells[4].text = txt
#表格宽度设置
widths = [3.99, 2.64, 3.15, 3.53, 4.11]
for i in range(0, 5):
    for cell in table1.columns[i].cells:
        cell.width = Cm(widths[i])
#表格中字体都设置为蓝色和居中
for i in range(0, 8):
    for j in range(0, 5):
        cell = table1.cell(i, j)
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.color.rgb = RGBColor(0x36, 0x5f, 0x91)
        cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
#表格填充颜色
for a,b in zip(range(1, 8), range(0, 7)):
    cell = table1.cell(a, 1)
    cell.paragraphs[0].runs[0].font.color.rgb = brush_color[b]
    cell.paragraphs[0].runs[0].font.bold = True

table1.style = 'ListCLF'
note = document.add_paragraph("注：（1）‘", style='N1')
fine = note.add_run("—")
fine.bold = True
fine.font.name = 'Arial'
fine.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
goon = document.paragraphs[-1].add_run("’为标准内，‘")
high = document.paragraphs[-1].add_run("↑")
high.bold = True
high.font.size = Pt(12)
high.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
document.paragraphs[-1].add_run("’为超出标准，‘")
medium = document.paragraphs[-1].add_run("↑")
medium.bold = True
medium.font.size = Pt(12)
medium.font.color.rgb = RGBColor(0xff, 0xff, 0x00)
document.paragraphs[-1].add_run("’为已达临界标准")
document.add_paragraph("（2）以上依据原始导出数据分析，见附件一或用户由平台自行导出", style='N1')
document.add_heading('1.2体检结论说明', level=2)
document.add_paragraph('此结论为%s变压器体检结果，' % TRANSFORMER +
                '变压器容量为%dkVA。' % KVA +
                '%s到%s共采集了%d次数据，高危共计%d次，隐患共%d次。' % (STARTDAY, ENDDAY, calc.freq, HIGHRISK, HIDDENRISK) +
                '以各体检项目在该时段内发生次数及超出各项标准次数进行分析得出：', style='Normal')
document.add_paragraph().add_run('电压体检结论：').bold = True
document.paragraphs[-1].add_run('该变压器在此数据时段内，正常工作时，电压数据%s' % conclusion[0] +
                '谐波电压含量约%d%%时段不符合国家标准，电网谐波电压含量%s' % ((1 - quality[1]) * 100, conclusion[1]))
document.add_paragraph().add_run('电流体检结论：').bold = True
document.paragraphs[-1].add_run('该变压器低压侧额定电流约%.1fA，' % (ele_n) +
                '电流%s' % conclusion[2] +
                '谐波电流%s' % conclusion[3] + '对应的最大值分别为%s' % calc.ithd_mv)
document.add_paragraph().add_run('功率因数体检结论：').bold = True
document.paragraphs[-1].add_run('工作时，功率因数在0.9~1之间的占比为%.1f%%，' % (quality[3] * 100) +
                '功率因数%s' % conclusion[4])
document.add_paragraph().add_run('负荷率体检结论：').bold = True
document.paragraphs[-1].add_run('该变压器容量为%dkVA，额定电流约%.1fA，' % (KVA, ele_n) +
                '工作时，变压器负荷率基本在%.1f%%左右运行，' % calc.lf_mean +
                '综合来看，监测期间变压器负荷率%s' % conclusion[5])
document.add_paragraph().add_run('三相电流不平衡体检结论：').bold = True
document.paragraphs[-1].add_run('正常工作时，三相电流不平衡度均在15%%以内的概率为%.1f%%，' % (quality[5] * 100) +
                '带负载时最大为%.1f%%，发生时间为%s。' % (calc.unb_max, calc.unb_maxtime) +
                '综合来看，三相电流不平衡度数据%s' % conclusion[6])
attention = document.add_paragraph(style='N1').add_run('注：分析数据来源为%s到%s，' % (STARTDAY, ENDDAY) +
                                 '共采集了%d次（正常情况15分钟为一个周期点，异常时即时推送）'% calc.freq)
attention.bold = True
attention.font.size = Pt(7.5)
document.add_page_break() #插入分页符
#第一章消耗0.07s

def explain(unquality):
    if unquality == '谐波电压':
        document.add_paragraph('电网中存在电压谐波主要是因为谐波电流，因为谐波电压是谐波电流流过线路阻抗时产生的，'
                               '对于特定的配电系统，谐波电流与谐波电压之间的关系为（欧姆定律）：谐波电压=谐波电流*电网阻抗。'
                               '电压谐波主要有以下设备产生：整流器、变频器、开关电源、静态换流器、'
                               '晶闸管系统及其他SCR控制系统等电力电子设备，和变压器、电动机、发电机等非线性设备，'
                               '以及电弧炉设备及气体电光源设备。', style='Normal')
        document.add_paragraph().add_run('可能会造成危害：').bold = True
        document.add_paragraph('1)	降低电网电压，增加线路损耗，浪费电网容量；\n\t'
                '2)	影响无功补偿设备，电容柜无法正常投切，严重时烧坏电容器，还会将电网谐波进一步放大；\n\t'
                '3)	影响设备的稳定性，尤其是对继电保护装置，容易在其动作时造成拉弧，危害特大。\n\t'
                '4)	造成电动机效率下降，噪声增大；\n\t'
                '5)	使低压开关设备产生误动作；\n\t'
                '6)	对通讯造成干扰，影响电力电子计量设备的准确性，甚至经常有电子烧坏现象。\n\t'
                '7)	增加电力变压器的铜损和铁损，影响变压器的使用容量、效率；造成变压器噪声增加，缩短变压器寿命。', style='Normal')
        document.add_paragraph().add_run('调整建议：').bold = True
        document.add_paragraph('电压谐波含量主要由谐波电流引起，'
                               '有效的治理电网谐波电流含量后电网谐波电压含量便会降低。', style='Normal')

    if unquality == '谐波电流':
        document.add_paragraph('供电系统谐波的定义是对周期性非正弦电量进行傅立叶级数分解，除了得到与电网基波频率相同的分量，'
                               '还得到一系列大于电网基波频率的分量，这部分电量称为谐波。', style='Normal')
        document.add_paragraph().add_run('可能会造成危害：').bold = True
        document.add_paragraph('1)	降低发电、输电及用电设备的效率，使线路过热甚至发生火灾。\n\t'
                '2)	谐波影响各种电气设备的正常工作，谐波会造成设备产生机械振动、噪声和过电压；\n\t'
                '3)	谐波使电容器、电缆等设备过热、绝缘老化、寿命缩短，以至损坏，发生火灾，甚至发生爆炸；\n\t'
                '4)	引起公用电网并联谐振和串联谐振，从而使谐波放大，使上述（1、2、3）的危害大大增加；\n\t'
                '5)	谐波会导致继电保护和自动装置的误动作，并会使电气测量仪表计量不准确；\n\t'
                '6)	谐波会对邻近的通信系统产生干扰，轻者产生噪声，降低通信质量；'
                '重者导致数据丢失，使通信系统无法正常工作。', style='Normal')
        document.add_paragraph().add_run('调整建议：').bold = True
        document.add_paragraph('1)	减少非线性用电设备与电源间的电气距离：减少系统阻抗，提高供电电压等级。\n\t'
                '2)	安装滤波器：目前对变电所侧和用户侧谐波治理的方法，多采用安装滤波器来减少'
                '谐波分量。滤波器分为有源滤波器和无源滤波器两大类。', style='Normal')

    if unquality == '功率因数':
        document.add_paragraph('功率因数就是有功功率与视在功率的比值，通俗地讲就是用电设备的实际出力'
                               '与用电设备的容量的比值，又简称为力率。', style='Normal')
        document.add_paragraph().add_run('可能会造成危害：').bold = True
        document.add_paragraph('1)	发电机组需要多做功，降低发电机的出力；\n\t'
                '2)	占用输、变、配电设备的资源，增加设备投资金额；\n\t'
                '3)	会使线路未端电压偏低，造成设备供电电压不足，会使线路电流、线损增加，增大电费支出；\n\t'
                '4)	力调电费罚款。', style='Normal')
        document.add_paragraph().add_run('调整建议：').bold = True
        document.add_paragraph('1)	检查无功补偿设备线路是否故障；\n\t'
                                '2)	电容本体烧坏、容量衰减，更换新电容本体；\n\t'
                                '3)	负载量增加，补偿容量不足，增加补偿容量；\n\t'
                                '4)	负载无功变化较快，传统补偿设备响应时间较慢，更换动态补偿设备；\n\t'
                                '5)	升级老旧无功补偿设备。', style='Normal')
    if unquality == '负荷率':
        document.add_paragraph('变压器的超负荷运转是指变压器运转时超越了铭牌上规则的电流值。变压器长时间处于严重超载'
                               '运行，不仅使变压器出力降低，而且还使变损增大，以及还会'
                               '出现变台设备损坏的情况，这不利于变压器经济安全运行。', style='Normal')
        document.add_paragraph().add_run('可能会造成危害：').bold = True
        document.add_paragraph('1)	输出电压降低；\n\t'
                '2)	变压器的损耗增大，寿命减少；\n\t'
                '3)	严重时烧坏变压器。', style='Normal')
        document.add_paragraph().add_run('调整建议：').bold = True
        document.add_paragraph('1)	更换变压器，增加变压器容量；\n\t'
                '2)	减少负载使用量；\n\t'
                '3)	长期监控变压器运行状态，合理分配负载。', style= 'Normal')
    if unquality == '三相电流不平衡度':
        document.add_paragraph('低压配电网络中，当三相电压对称三相负荷也对称的情况下，三相电流也将是对称的。'
                               '各相负荷的有功功率和无功功率不全相等且三相电流的相量和不等于零，'
                               '称为三相电流不平衡。', style='Normal')
        document.add_paragraph().add_run('可能会造成危害：').bold = True
        document.add_paragraph('1)	增加线路和变压器的电能损耗；\n\t'                        
                '2)	配变出力减少；\n\t'
                '3)	产生零序电流；\n\t'
                '4)	影响用电设备的安全运行；\n\t'
                '5)	电动机效率降低。', style='Normal')
        document.add_paragraph().add_run('调整建议：').bold = True
        document.add_paragraph('1)	均匀合理的分配单相负载；\n\t'
                '2)	调整负荷较大一相至其他两相；\n\t'
                '3)	安装三相电流不平衡调节装置。', style='Normal')

if '↑' in result:
    document.add_heading('1.3体检不正常项目说明', level=2)
    sum = 0
    for a, b in zip(result, chapter):
        if a == '↑' and b != '电压' and b != '电流':
            sum += 1
            print(b)
            document.add_heading('1.3.' + str(sum) + b, level=3)
            explain(b)
    document.add_page_break() #插入分页符
document.add_heading('二、专项体检数据分析说明', level=1)

for a, b in zip(result, chapter):
    if a == '↑':
        risk_list.append(b)

document.add_paragraph('数据点为%s变压器，变压器容量%dkVA。'% (TRANSFORMER, KVA) +
            '从%s到%s共采集了%d次数据，'%(STARTDAY, ENDDAY, calc.freq) +
            '结合各项标准，此时段内存在%s超标隐患。' % ('、'.join(risk_list)) +
            '此时段电压数据超出标准限值约%d次；谐波电压含量超出标准限值约%d次；' % (riskamount[0], riskamount[1]) +
            '电流超出变压器额定电流值约%d次；功率因数超出标准限值约%d次；' % (riskamount[2], riskamount[3]) +
            '变压器负荷率85%%以上约%d次；三相电流不平衡超出标准限值%d次。'% (riskamount[4], riskamount[5]), style='Normal')
document.add_heading('2.1电压数据', level=2)
document.add_heading('2.1.1 电压数据体检分析', level=3)
document.add_paragraph('根据国家标准《GB/T 12325-2008》中规定单相220V供电电压允许偏差为标称系统电压的+7%%、-10%%。'
            '由此计算出电压标准上限值为235.4V。'
            '从以下分析中可以看出A、B、C三相电压达标率为%.1f%%左右。' % (qramount[0] * 100), style='Normal')
document.add_heading('1）电压趋势图', level=4)
document.add_paragraph('%s至%s时段内监测数据，标准限值198V~235V，' % (STARTDAY, ENDDAY) +
            '下图中最大值%.1f，最小值%.1f，用户可在平台内自行查看，' % (max_trend[0], min_trend[0]) +
            '一周内每日显示96个数据点，一月内显示每日平均值，'
            '可根据此图(该图使用2小时均值聚合而成)看出该时段内电压变化状态，用来分析用电情况。', style='Normal')
document.add_picture('./pic/U_trend.png', height=Cm(9.2), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74) #五号一个字符长度0.371厘米 缩进了两个字符
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-1  电压趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_page_break() #插入分页符
document.add_heading('2）电压数据概率分布', level=4)
document.add_picture('./pic/U.png', height=Cm(6.6), width=Cm(17.2)) #电压图17.2*6.6
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-2 电压值概率分布图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('3）电压数据合格率', level=4)
document.add_paragraph('取导出数据中各相电压实际超出标准235.4V次数，与总次数对比分析得出该合格率，'
                       '此分析结果可用于查看该时段内各相电压在标准范围的占比，来判断该时段内的电压是否合格。', style='Normal')
document.add_picture('./pic/Upie.png', height=Cm(5.71), width=Cm(17.21)) #环形图17.1*5.71
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-3 电压合格率环形图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.1.2 电压释义', level=3)
document.add_heading('1）电压的定义', level=4)
document.add_paragraph('电压（voltage），也称作电势差或电位差，是衡量单位电荷在静电场中由于电势不同所产生的能量差的物理量。'
                       '其大小等于单位正电荷因受电场力作用从A点移动到B点所做的功，'
                       '电压的方向规定为从高电位指向低电位的方向。', style='Normal')
document.add_page_break()
document.add_heading('2）电压的算法', level=4)
p = document.add_paragraph('线电压算法:', style='Normal')
r = p.add_run()
r.add_picture('./pic/U_line.png', width=Cm(3.36), height=Cm(1.19))
p = document.add_paragraph('相电压算法:', style='Normal')
r = p.add_run()
r.add_picture('./pic/U_phase.png', width=Cm(1.80), height=Cm(1.19))
document.add_heading('3）参考标准', level=4)
document.add_paragraph('参考《GB/T 12325-2003、2008》电能质量供电电压偏差，'
                       '单相220V供电电压允许偏差为标称系统电压的+7%、-10%。', style='Normal')
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.2 谐波电压数据', level=2)
document.add_heading('2.2.1 谐波电压体检数据分析小结', level=3)
document.add_paragraph('国家标准《GB/T14549-1993》中相电压最大谐波畸变率限值为5%%，从以下分析中可以看出A、B、C各相电压总谐'
                       '波畸变率已明显超出国家5%%标准，最大值已达%.1f%%，' % max_value[1] +
                       '电网电压谐波含量需要治理。', style='Normal')
document.add_heading('1）谐波电压趋势图', level=4)
document.add_paragraph('%s至%s时段内监测数据，标准限值5%%，' %(STARTDAY, ENDDAY) +
                       '下图中最大值%.1f%%，最小值%.1f%%，用户可在平台内自行查看，' %(max_trend[1], min_trend[1]) +
                       '一周内每日显示96个数据点，一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)看出该时段内谐波电压变'
                       '化状态，用来分析谐波电压发生在哪个时段内。', style='Normal')
document.add_picture('./pic/UTHD_trend.png', height=Cm(9.2), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-4  谐波电压含量趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('2）谐波电压数据详情', level=4)
document.add_paragraph('根据导出数据中各相电压谐波含量发生次数分析得出该分布图，标准限值5%。此分析结果可用于观察该时段内电压谐波'
                       '含量主要分布在多少范围内，并可以看出该范围内谐波电压含量发生的概率，以及超出标准的概率。', style='Normal')
document.add_picture('./pic/U_THD.png', height=Cm(6.6), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-5  谐波电压含量概率分布图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('3）谐波电压含量合格率', level=4)
document.add_paragraph('取导出数据中各相谐波电压含量实际超出5%标准次数，与总次数对比分析得出该合格率，此分析结果可用于查看该时段内各'
                       '相谐波电压含量在标准范围的占比，来判断该时段内的电压是否合格。', style='Normal')
document.add_picture('./pic/UTHD_pie.png', height=Cm(5.74), width=Cm(17.21))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-6 谐波电压含量合格率环形图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_page_break()
document.add_heading('2.2.2 谐波电压释义', level=3)
document.add_heading('1）谐波电压定义', level=4)
document.add_paragraph('谐波电压是谐波电流流过线路阻抗时产生的，对于特定的配电系统，谐波电流与谐波电压之间的关系如下（欧姆定律）：'
                       '谐波电压=谐波电流*电网阻抗。', style='Normal')
document.add_heading('2）电压总谐波含量算法', level=4)
document.add_picture('./pic/UTHD_percent.png', height=Cm(1.2), width=Cm(3.38))
document.add_paragraph('式中：Uh——第h次谐波电压（方根均值）；\n\t\t  U1——基波电压（方根均值）；', style='Normal')
document.add_heading('3）参考标准', level=4)
document.add_paragraph('依据《GB/T14549-1993》电能质量公用电网谐波国家标准进行。\n\t'
                       '国标《GB/T14549-1993》中公用电网谐波电压（相电压）谐波含量限值：', style='Normal')
table2 = document.add_table(rows=1, cols=4)
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = '电网标称电压\t(kV)'
hdr_cells2[1].text = '电压总谐波畸变率(%)'
hdr_cells2[2].text = '各次谐波电压含有率(%)'
for i in range(0, 3):
    hdr_cells2[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells2[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)

contents2 = (('奇次', '偶次'),
           ('0.38', '5.0', '4.0', '2.0'))
addr = table2.add_row().cells
addr[2].text = contents2[0][0]
addr[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
addr[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
addr[3].text = contents2[0][1]
addr[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
addr[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
addr1 = table2.add_row().cells
for j in range(0, 4):
    addr1[j].text = contents2[1][j]
    addr1[j].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr1[j].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table2widths = [3.07, 3.07, 2.23, 2.23]
for i in range(0, 4):
    for cell in table2.columns[i].cells:
        cell.width = Cm(table2widths[i])

a = table2.cell(0, 0)
b = table2.cell(1, 0)
c = table2.cell(0, 1)
d = table2.cell(1, 1)
e = table2.cell(0, 2)
f = table2.cell(0, 3)
a.merge(b)
c.merge(d)
e.merge(f)
table2.style = 'ListCLF2'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
document.add_paragraph('表2-1 《GB/T14549-1993》中公用电网谐波电压（相电压）谐波含量限值')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.3 电流数据', level=2)
document.add_heading('2.3.1 电流数据体检分析小结', level=3)
document.add_paragraph('该变压器容量为%dkVA，则二次侧额定电流约为%dA。' % (KVA, ele_n) +
            '监测时间段三相电流%s，' % conclusion[4] +
            '单相最大电流为%.1fA，' % max_value[2] +
            '最高时超出额定电流%.1f%%左右。' % (((max_value[2] / ele_n) -1) * 100) , style='Normal')
document.add_page_break()
document.add_heading('1）电流趋势图', level=4)
document.add_paragraph('%s至%s时段内监测数据，标准限值%dA（变压器额定电流），' % (STARTDAY, ENDDAY, KVA) +
                       '下图中最大值%.1fA，最小值%.1fA，用户可在' % (max_trend[2], min_trend[2]) +
                       '平台内自行查看，一周内每日显示96个数据点，一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)'
                       '看出该时段内电流变化状态，用来分析用电情况。', style='Normal')
document.add_picture('./pic/I_trend.png', height=Cm(9.2), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-7 电流趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('2）电流数据详情', level=4)
document.add_paragraph('根据导出数据中各相电流发生次数分析得出该概率图，标准限值%dA。此分析结果可用于观察该时段内电流主要分' % ele_n +
                       '布在多少范围内，并可以看出该范围内电流值发生的概率，以及超出标准的概率。', style='Normal')
document.add_picture('./pic/I.png', height=Cm(6.6), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-8 电流概率分布图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_page_break()
document.add_heading('3）电流数据合格率', level=4)
document.add_paragraph('取导出数据中各相电流实际超出二次额定电流的次数，与总次数对比分析得出该合格率，此分析结果可用于查看该时段内各相电'
                       '流在标准范围内、外的占比，来判断该时段内的电流数据是否合格。', style='Normal')
document.add_picture('./pic/I_Pie.png', height=Cm(5.74), width=Cm(17.21))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-9 电流合格率环形图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.3.2 电流释义',level=3)
document.add_heading('1）电流定义', level=4)
document.add_paragraph('单位时间里通过导体任一横截面的电量叫做电流强度，简称电流。', style='Normal')
document.add_heading('2）电流算法', level=4)
r0 = document.add_paragraph('电流算法：', style='Normal').add_run()
r0.add_picture('./pic/I_format.png', width=Cm(3.35), height=Cm(1.27))
document.add_heading('3) 参考标准', level=4)
document.add_paragraph('根据国标JGJ16-2008《民用建筑电气设计手册》中4.3.2要求：配变压器长期工作电流不宜大于变压器额'
                       '定电流的85%。', style='Normal')
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.4 谐波电流数据', level=2)
document.add_heading('2.4.1 谐波电流体检数据分析小结', level=3)
document.add_paragraph('从以下分析中可以看出三相各分次谐波电流值，根据‘国标《GB/T14549-1993》注入公共连接点的谐波电流允许值’中'
                       '规定，该电网谐波电流第%s次均有超出国家标准范围，' % talk.strs +
                       '谐波电流%s' % conclusion[3] + '对应的最大值分别为%s' % calc.ithd_mv, style='Normal')
document.add_page_break()
document.add_heading('1）谐波电流趋势图', level=4)
document.add_paragraph('12月份时段内监测数据，'
                       '下图中最大值%.1fA，最小值%.1fA，' % (max_trend[3], min_trend[3]) +
                       '用户可在平台内自行查看，一周内每日显示96个数据点，一月内显示每日平均值，'
                       '可根据此图(该图使用2小时均值聚合而成)看出该时段内谐波电流变化状态，'
                       '用来分析电网谐波含量情况。', style='Normal')
document.add_picture('./pic/ITHD_trend.png', height=Cm(8.47), width=Cm(16.36))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-10 总谐波电流值趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('2）谐波电流数据详情', level=4)
document.add_paragraph('取导出数据中每相各分次电流谐波最大值与平均值分析得出如下两图。此分析结果可用于观察该时段'
                       '内各分次谐波电流值，主要谐波电流为几次，并可以根据国标《GB/T14549-1993》看出该'
                       '范围内各分次谐波电流值是否超出标准，以判断谐波电流是否合格。', style='Normal')
document.add_picture('./pic/HVmax.png', height=Cm(7.38), width=Cm(16.17))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-11 各分次谐波电流最大值')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_page_break()
document.add_picture('./pic/HVmean.png', height=Cm(7.4), width=Cm(16.7))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-12 各分次谐波电流平均值')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('各相各分次谐波电流值合格率', level=4)
document.add_paragraph('取导出数据中各相电流分次谐波实际超出标准次数，与总次数对比分析得出该合格率，此分析结果可用于查看该时段'
                       '内各相电流谐波含量在标准范围的占比，来判断该时段内的电流谐波是否合格。', style='Normal')
document.add_picture('./pic/HVqrate.png', height=Cm(7.4), width=Cm(16.7))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-13 各分次谐波电流合格率图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_page_break()
document.add_heading('2.4.2 谐波电流释义', level=3)
document.add_heading('1）谐波电流定义', level=4)
document.add_paragraph('供电系统谐波的定义是对周期性非正弦电量进行傅立叶级数分解，除了得到与电网基波频率相同的分量，还得到一系列'
                       '大于电网基波频率的分量，这部分电量称为谐波。', style='Normal')
document.add_heading('2）电流总谐波含量算法', level=4)
r1 = document.add_paragraph(style='Normal').add_run()
r1.add_picture('./pic/thdi_format.png', width=Cm(3.3), height=Cm(1.2))
document.add_heading('3）参考标准', level=4)
document.add_paragraph('依据《GB/T14549-1993》电能质量公用电网谐波国家标准进行。\n\t'
                       '国标《GB/T14549-1993》中注入公共连接点的谐波电流允许值：', style='Normal')
table3 = document.add_table(rows=1, cols=14)
hdr_cells3 = table3.rows[0].cells
thead3 = ('标准电压（KV）', '基准短路容量（MVA）', '谐波次数及谐波电流允许值（A）')
hdr_cells3[0].text = thead3[0]
hdr_cells3[1].text = thead3[1]
hdr_cells3[2].text = thead3[2]
for i in range(0, 3):
    hdr_cells3[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells3[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
tr3_1 = ('', '', '3', '5', '7', '9', '11', '13', '15', '17', '19', '21', '23', '25')
tr3_2 = ('0.38', '10', '62', '62', '44', '21', '28', '24', '12', '18', '16', '8.9', '14', '12')
table3_add1 = table3.add_row().cells
for i in range(2, 14):
    table3_add1[i].text = tr3_1[i]
    table3_add1[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table3_add1[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table3_add2 = table3.add_row().cells
for i in range(0, 14):
    table3_add2[i].text = tr3_2[i]
    table3_add2[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table3_add2[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
g = table3.cell(0, 2)
h = table3.cell(0, 13)
i = table3.cell(0, 0)
j = table3.cell(1, 0)
k = table3.cell(0, 1)
l = table3.cell(1, 1)
g.merge(h)
i.merge(j)
k.merge(l)
for i in range(0, 2):
    for cell in table3.columns[i].cells:
        cell.width = Cm(2.8)
for i in range(2, 14):
    for cell in table3.columns[i].cells:
        cell.width = Cm(1.0)
table3.style = 'ListCLF2'
document.add_paragraph('表2-2 《GB/T14549-1993》中注入公共连接点的谐波电流允许值')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r2 = document.add_paragraph('由于PCC点的短路容量不同于假定基准最小短路容量,应按照国标附录B进行换'
                            '算,换算公式如下：\n\t\t', style='Normal').add_run()
r2.add_picture('./pic/thdi_trans.png', height=Cm(1.25), width=Cm(2.11))
document.add_paragraph('式中∶Sk1:公共连接点的最小短路容量,MVA；\n\t\t  Sk2:基准短路容量,MVA；\n\t\t  Ihp:表2中的第h次谐波'
                       '电流允许值,A；\n\t\t  Ih:短路容量为Sk1时的第h次谐波电流允许值,A。', style='Normal')
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.5 功率因数数据', level=2)
document.add_heading('2.5.1 总功率因数体检数据分析小结', level=3)
document.add_paragraph('根据国标《JGJ16-2008》中要求，功率因数不宜低于0.9。'
                       '从以下分析中可以看出功率因数0.9到1的占比为%.1f%%，' % (qramount[3] * 100) +
                       '0.9以下占比%.1f%%，安全提示：若是功率因数偏低，应加强对功率因数' % ((1 - qramount[3]) * 100) +
                       '的巡查监管，避免可能带来的功率因数罚款。', style='Normal')
document.add_page_break()
document.add_heading('1）总功率因数趋势图', level=4)
document.add_paragraph('12月份时段内监测数据，标准限值0.9，'
                       '下图中最大值%.1f，最小值%.1f，用户可在平台内自行查看，' % (max_trend[4], min_trend[4]) +
                       '一周内每日显示96个数据点，一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)看出该时段内功率因数变'
                       '化状态，用来分析功率因数是否达标。')
document.add_picture('./pic/PF_trend.png', height=Cm(8.7), width=Cm(16.3))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-14 总功率因数趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('2）总功率因数体检数据详情', level=4)
document.add_paragraph('根据导出数据中功率因数发生次数和该时段内超出0.9次数与总次数的对比进行分析。此分析结果可用于观察该时段内功'
                    '率因数主要分布在多少范围内，并可以看出该范围内功率因数发生的概率，以及超出标准的概率，'
                    '可用于分析该时段内功率因数在标准范围内、外的占比，来判断该时段内的功率因数是否合格。', style='Normal')
document.add_picture('./pic/PF.png', height=Cm(6.2), width=Cm(17.15))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-15 总功率因数分析图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_page_break()
document.add_heading('2.5.2功率因数释义', level=3)
document.add_heading('1）功率因数定义', level=4)
document.add_paragraph('所谓功率因数就是有功功率与视在功率的比值，通俗地讲就是用'
                       '电设备的实际出力与用电设备的容量的比值，又简称为力率。', style='Normal')
document.add_heading('2）功率因数的算法', level=4)
document.add_paragraph('在交流电路中，电压与电流之间的相位差(Φ)的余弦叫做功率因数，用符号cosΦ表示，在数值上，功率因数是有功功率和'
                       '视在功率的比值，即cosΦ=P/S，功率因数表是指测量交流电路中有功功率'
                       '与视在功率的比值或测量电压、电流间相位角余弦的电表；', style='Normal')
document.add_picture('./pic/PF_format.png', height=Cm(2.54), width=Cm(8.9))
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('3）参考标准', level=4)
document.add_paragraph('（1）国标《JGJ16-2008》民用建筑电气设计规范中第3.6.2条规定：10（6）kV及以下无功补偿宜在配电变压器低压侧'
                       '集中补偿，且功率因数不宜低于0.9。高压侧的功率因数指示应符合当地供电部门的规定。'
                       '\n\t（2）现行的《国家电网公司电力系统电压质量和无功电力管理规定》规定，100kVA及以上10kV供电的电力用户在'
                       '用户在用户高峰负荷时变压器高压侧功率因数不宜低于0.95；其他电力用户功率因数不宜低于0.9。', style='Normal')
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.6 负荷率数据', level=2)
document.add_heading('2.6.1 负荷率体检数据分析小结', level=3)
document.add_paragraph('在正常工作时，最大负荷率在%.1f%%，' % max_value[5] +
                       '负荷率%s' % conclusion[5] +
                       '安全提示:若是负荷率长期在过载状态工作可能会发生安全隐患。', style='Normal')
document.add_heading('1）负荷率趋势图', level=4)
document.add_paragraph('%s至%s时段内监测数据，标准限值85%%，' % (STARTDAY, ENDDAY) +
                       '下图中最大值%.1f%%，最小值%.1f%%，' % (max_trend[5], min_trend[5]) +
                       '用户可在平台内自行查看，一周内每日'
                       '显示96个数据点，一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)看'
                       '出该时段内负荷率变化状态，用来分析用电情况。', style='Normal')
document.add_page_break()
document.add_picture('./pic/LF_trend.png',  height=Cm(9.1), width=Cm(17.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-16 变压器负荷率趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('2）负荷率数据详情', level=4)
document.add_paragraph('根据导出数据中负荷率发生次数和该时段内超出85%次数与总次数的对比进行分析。此分析结果可用于观察该时'
                       '段内变压器负荷率主要分布在多少范围，并可以看出该范围内负荷率发生的概率，以及超出标准的概率，可用于分析该'
                       '时段内负荷率在标准范围内、外的占比，来判断该时段内的负荷率是否合格。', style='Normal')
document.add_picture('./pic/LF.png', height=Cm(6.3), width=Cm(17.0))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-17 负荷率数据分析图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.6.2 变压器负荷率释义', level=3)
document.add_heading('1）负荷率定义', level=4)
document.add_paragraph('一定时间内，变压器平均输出的视在功率与变压器额定容量之比称之为负荷率。', style='Normal')
document.add_heading('2）负荷率的算法', level=4)
r3 = document.add_paragraph('负荷率 = S/变压器容量 * 100%；\n\t视在功率', style='Normal').add_run()
r3.add_picture('./pic/S_format.png', height=Cm(0.69), width=Cm(2.76))
document.add_paragraph('总有功功率P = U*I*cosφ\n\t总无功功率Q = U*I*sinφ', style='Normal')
document.add_heading('3）参考标准', level=4)
document.add_paragraph('根据国标JGJ16-2008《民用建筑电气设计手册》中4.3.2要求：配电变压器的长期工'
                       '作负载率不宜大于85%。', style='Normal')
document.add_heading('2.7 三相电流不平衡数据', level=2)
document.add_heading('2.7.1三相电流不平衡体检数据分析小结', level=3)
document.add_paragraph('正常工作时三相电流不平衡度最大为%.1f%%，' % max_value[6] +
                       '其中已经剔除了轻负载时的不平衡度，可以得出，三相电流不平衡度%s' , style='Normal')
document.add_heading('1）三相电流不平衡趋势图', level=4)
document.add_paragraph('%s至%s时段内监测数据，标准限值15%%，' % (STARTDAY, ENDDAY) +
                       '下图中最大值%.1f%%，最小值%.1f%%，用户可在平台内自行查看，' % (max_trend[6], min_trend[6])+
                       '一周内每日显示96个数据点，一月内显示每日平均值，可根据此图(该图使用2小时均值聚合而成)'
                       '看出该时段内三相电流不平衡状态，用来分析各相用电情况。', style='Normal')
document.add_picture('./pic/Unb_trend.png', height=Cm(8.8), width=Cm(16.2))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-18 三相电流不平衡趋势图')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_page_break()
document.add_heading('3）三相电流不平衡数据详情', level=4)
document.add_paragraph('根据导出数据中三相不平衡度发生次数和该时段内超出15%次数与总次数的对比进行分析。此分析结果可用于观察该'
                       '时段内变压器三相电流平衡度主要分布在多少范围，并可以看出该范围内发生的概率，以及超出标准的概率，'
                       '可用于分析该时段内三相不平衡在标准范围内、外的占比，来判断该时段内的三相电流平衡度是否合格。', style='Normal')
document.add_picture('./pic/Unb.png', height=Cm(6.57), width=Cm(16.9))
document.paragraphs[-1].paragraph_format.left_indent = -Cm(0.74)
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('图2-19 三相电流不平衡数据分析')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_heading('2.7.2三相不平衡释义', level=3)
document.add_heading('1)变压器三相不平衡定义', level=4)
document.add_paragraph('是指在电力系统中三相电流（或电压）幅值不一致，且幅值差超过规定范围。', style='Normal')
document.add_heading('2)不平衡度的计算公式', level=4)
document.add_picture('./pic/Unb_format.png', height=Cm(1.20), width=Cm(9.63))
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('3)参考标准', level=4)
document.add_paragraph('国标《GB/T 1094-2013》下《SD292-1988》第五节 负荷管理 第7.5.2条  变压器的三相负荷力求平衡，不平'
                       '衡度不应大于15%，只带少量单相负荷的三相变压器，零线电流不应超过'
                       '额定电流的25%，不符合上述规定时，应将负荷进行调整。', style='Normal')
document.add_paragraph('--------------------------------------------------------------------', style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
document.add_page_break()
document.add_heading('三、我们的服务', level=1)
document.add_heading('3.1健康体检', level=2)
document.add_paragraph('在%s至%s时段内，为您的%s变压器' % (STARTDAY, ENDDAY, TRANSFORMER) +
                       '做了电流数据、电压数据、谐波电压、谐波电流、三相不平衡、功率因数、变压器负'
                       '荷率数据的采集及分析管理服务，数据界面在PC和微信公众号界面均有呈现，采集时间为15分钟一次'
                       '（异常时即时采集），可随时随地读取、查看各个回路参数信息。电能健康体检设备可以动态体检，'
                       '把每天采集的多个点形成趋势图记录下来，以便后续分析查看。可监测多个项目，如您需要监测'
                       '其他参数，我们可以根据您的需求提供相应的监测服务。', style='Normal')
document.add_heading('3.2隐患管理', level=2)
document.add_paragraph('设备采集的数据经过阿里云混合云服务器的解析，结合国家标准对隐患级别进行分类、分级管理。使得其各类'
                       '参数的更清楚、更明确的呈现出来，隐患变得一目了然。\n\t在数据时段内，'
                       '对各监测项%d条信息进行了' % (HIGHRISK + HIGHRISK) +
                       '分析，其中高危%d条，隐患%d条，详情如下表格：' % (HIGHRISK, HIDDENRISK)) #手动填写
table4 = document.add_table(rows=1, cols=9)
th4 = ('体检项', '电压', '电流', '电压谐波含量', '谐波电流', '电流不平衡', '功率因数', '负荷率', '温度')
tr4 = (('高危（次）',	0, 255, 2, 2, 0, 0, 79, '-'), ('隐患（次）', 527, 735, 5, 5, 302, 2, 721, '-'),
       ('总计', 527, 990, 7,	7, 302, 2, 800, '-'))
hdr_cells4 = table4.rows[0].cells
for i in range(0, 9):
    hdr_cells4[i].text = th4[i]
    hdr_cells4[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells4[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)

for i in range(0,3):
    table4_add = table4.add_row().cells
    for j in range(0, 9):
        table4_add[j].text = str(tr4[i][j])
        table4_add[j].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table4_add[j].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)

table4.style = 'ListCLF3'
table4widths = (1.75, 1.75, 1.6, 3, 2.25, 2.4, 2.25, 1.75, 1.25)
for i in range(0, 9):
    for cell in table4.columns[i].cells:
        cell.width = Cm(table4widths[i])
document.add_heading('3.3健康干预', level=2)
document.add_heading('\t“未病先防，既病防变”', level=4)
document.add_paragraph('在隐患分类分级后，健康干预也会针对不同隐患等级有不同的推送方式，这些通知方式也是可供用户自主选择的，在'
                       '用户没有提出通知方式之前，我们对用户的健康干预考虑了有效实时，不会对用户生活造成干扰。当预警达到一定的数量'
                       '的时候，而这时用户自身没有留意到风险本身，%s的专业团队' % NANDECLOUD_BRAND +
                       '便会对风险用户进行致电，甚至就风险问题根源'
                       '进行上门排查、解决。\n\t在数据时段内，系统共监测到%d条预警信息，'
                       '并以线上和客服方式通知，共预警了XXX条，详情如下表格：', style='Normal')
table5 = document.add_table(rows=1, cols=5)
th5 = ('干预方式', '微信通知', '短信通知', '邮件通知', '客服致电')
tr5 = ('干预次数', 388, 0, 0, 0)
hdr_cells5 = table5.rows[0].cells
for i in range(0, 5):
    hdr_cells5[i].text = th5[i]
    hdr_cells5[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells5[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table5_add = table5.add_row().cells
for j in range(0, 5):
    table5_add[j].text = str(tr5[j])
    table5_add[j].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table5_add[j].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table5.style = 'ListCLF3'
document.add_page_break()
document.add_heading('3.4 O2O服务', level=2)
document.add_paragraph('线上对线上，当用户在有困惑时，点击“呼叫服务”，%s就能知晓客户有疑问与反馈，在客户未' % NANDECLOUD_BRAND +
                        '拨通客服电话，乃至在拨通电话前，%s的健康指导师会直接联系客户，或微信，' % NANDECLOUD_BRAND +
                        '或电话，直至将客户的心中的疑虑排除。\n\t已为您提供O2O服务0次。', style='Normal')
r4 = document.add_paragraph().add_run()
r4.add_picture('./pic/phone_poweryun.png', height=Cm(4.15), width=Cm(6.59))
r4.add_picture('./pic/service_poweryun.png', height=Cm(10.37), width=Cm(5.84))
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('3.5电能健康档案', level=2)
document.add_paragraph('客户使用我们的%s，自设备建点那一刻起就建立了档案，' % NANDECLOUD_BRAND +
                       '您的电能健康档案已于XXXX年XX月XX日建立，'
                       '并长期有效存储在阿里云服务器中，数据传输以15分钟为间隔，都存档在案，若追溯，最长可长达2年。'
                       '也就是病史可查，若经过治理，治理前与治理后的数据可轻松调用对比。', style='Normal')
document.add_paragraph('如您对电能质量隐患有治理需求可与我公司联系，我公司将派专业服务团队进行对接！', style='N1')
r5 = document.add_paragraph().add_run()
r5.add_picture('./pic/poweryun_seal.png', height=Cm(2.31), width=Cm(3.16))
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
document.add_paragraph('技术中心技术二部 %s' % (datetime.datetime.strftime(now, '%y-%m-%d')), style='Normal')
document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT



document.save('./text/test.docx')
elapsed = (time.clock() - start)
print("Time used:", elapsed)
