# -*- encoding: utf-8 -*-
"""
@File    : autodetection.py
@Time    : 2019-7-4 上午 10:20
@Author  : major
@Email   : major3428@foxmail.com
@Software: PyCharm
"""

from docx import Document
from docx.shared import Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import numpy as np

from docx.oxml.ns import qn
import datetime
from detectconf import *
import os
import pandas as pd

# 导入数据
df = pd.read_excel(FILENAME, index_col=None)
now = datetime.datetime.now()
docx = Document(MODELPATH)

company = docx.paragraphs[1].add_run(text=COMPANY)
company.font.size = Pt(36)
company.font.bold = True
sta = docx.paragraphs[7].add_run(text=STAFF)
sta.font.size = Pt(22)
sta.font.name = u'微软雅黑'
wri = docx.paragraphs[8].add_run(text=STAFF)
wri.font.size = Pt(22)
wri.font.name = u'微软雅黑'
dat = docx.paragraphs[13].add_run(text=now.strftime('%Y')+'年 '+now.strftime('%m')+'月')
dat.font.size = Pt(18)
dat.font.name = 'Microsoft YaHei UI'
last = docx.paragraphs[-1].clear()
last.add_run(text='1 测试概况')
records = (('项目名称', COMPANY + '测试报告'),
           ('项目地址', 'XXX'),
           ('联系人联系方式', 'XXX 1XXXXX'),
           ('变压器编号及容量', 'XXXkVA 箱变'),
           ('负荷类型', 'XXX'),
           ('基本情况描述', '1、XXX'),
           ('', '2、XXX'),
           ('', '3、XXX'),
           ('测试时间', 'XXX')
           )
table1 = docx.add_table(rows=0, cols=2)
for state, content in records:
    row_cells = table1.add_row().cells
    row_cells[0].paragraphs[0].add_run(state).font.name = 'Microsoft YaHei UI'
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = Cm(-0.74)
    row_cells[1].paragraphs[0].add_run(content).font.name = 'Microsoft YaHei UI'
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = Cm(-0.74)
table1.cell(5, 0).merge(table1.cell(7, 0))
table1.cell(5, 1).merge(table1.cell(7, 1))
#表格宽度设置
widths = [4.01, 11.02]
for i in range(0, 2):
    for cell in table1.columns[i].cells:
        cell.width = Cm(widths[i])
table1.style = 'Table Grid'
docx.add_page_break()
docx.add_heading('2 测试依据', level=1)
docx.add_paragraph('该项测试依据GB/T14549-93电能质量 公用电网谐波国家标准进行。')
docx.add_paragraph('GB/T14549-93各级电压等级谐波限值规定如下表1, 公共连接点的全部用户向该点注入的谐波电流允许值见表2。')
list1 = docx.add_paragraph().add_run('表1：公用电网谐波电压（线电压）限值')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
list1.font.bold = True
table2 = docx.add_table(rows=1, cols=4)
hdr_cells2 = table2.rows[0].cells
harmGb = ('电网标称电压\t(kV)', '电压总谐波畸变率(%)', '各次谐波电压含有率(%)')
for i in range(0, 3):
    hdr_cells2[i].text = harmGb[i]
    hdr_cells2[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells2[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
contents2 = (('奇次', '偶次'),
           ('0.38', '5.0', '4.0', '2.0'))
addr = table2.add_row().cells
addr[2].text = contents2[0][0]
addr[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
addr[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
addr[3].text = contents2[0][1]
addr[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
addr[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
addr1 = table2.add_row().cells
for j in range(0, 4):
    addr1[j].text = contents2[1][j]
    addr1[j].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    addr1[j].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table2widths = [3.76, 4.28, 3.23, 3.23]
for i in range(0, 4):
    for cell in table2.columns[i].cells:
        cell.width = Cm(table2widths[i])

a = table2.cell(0, 0)
b = table2.cell(1, 0)
c = table2.cell(0, 1)
d = table2.cell(1, 1)
e = table2.cell(0, 2)
f = table2.cell(0, 3)
a.merge(b)
c.merge(d)
e.merge(f)
table2.style = 'MyTable'
styles = [s for s in docx.styles if s.type == WD_STYLE_TYPE.TABLE]
for style in styles:
    print(style.name)
list2 = docx.add_paragraph().add_run('表2：公用电网谐波电压（线电压）限值')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
list2.font.bold = True
table3 = docx.add_table(rows=1, cols=14)
hdr_cells3 = table3.rows[0].cells
thead3 = ('标准电压（KV）', '基准短\n\t路容量\n\tMVA', '谐波次数及谐波电流允许值（A）')
for i in range(0, 3):
    hdr_cells3[i].text = thead3[i]
    hdr_cells3[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells3[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
tr3_1 = ('', '', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12', '13')
tr3_2 = ('0.38', '10', '78', '62', '39', '62', '26', '44', '19', '21', '16', '28', '13', '24')
table3_add1 = table3.add_row().cells
for i in range(2, 14):
    table3_add1[i].text = tr3_1[i]
    table3_add1[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table3_add1[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table3_add2 = table3.add_row().cells
for i in range(0, 14):
    table3_add2[i].text = tr3_2[i]
    table3_add2[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table3_add2[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
g = table3.cell(0, 2)
h = table3.cell(0, 13)
i = table3.cell(0, 0)
j = table3.cell(1, 0)
k = table3.cell(0, 1)
l = table3.cell(1, 1)
g.merge(h)
i.merge(j)
k.merge(l)
for i in range(0, 2):
    for cell in table3.columns[i].cells:
        cell.width = Cm(2.9)
for i in range(2, 14):
    for cell in table3.columns[i].cells:
        cell.width = Cm(1.0)
table3.style = 'MyTable'
docx.add_paragraph()

table4 = docx.add_table(rows=1, cols=14)
hdr_cells4 = table4.rows[0].cells
thead4 = ('标准电压（KV）', '基准短\n\t路容量\n\tMVA', '谐波次数及谐波电流允许值（A）')
for i in range(0, 3):
    hdr_cells4[i].text = thead4[i]
    hdr_cells4[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells4[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
tr4_1 = ('', '', '14', '15', '16', '17', '18', '19', '20', '21', '22', '23', '24', '25')
tr4_2 = ('0.38', '10', '11', '12', '9.7', '18', '8.6', '16', '7.8', '8.9', '7.1', '14', '6.5', '12')
table4_add1 = table4.add_row().cells
for i in range(2, 14):
    table4_add1[i].text = tr4_1[i]
    table4_add1[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table4_add1[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table4_add2 = table4.add_row().cells
for i in range(0, 14):
    table4_add2[i].text = tr4_2[i]
    table4_add2[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table4_add2[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
g = table4.cell(0, 2)
h = table4.cell(0, 13)
i = table4.cell(0, 0)
j = table4.cell(1, 0)
k = table4.cell(0, 1)
l = table4.cell(1, 1)
g.merge(h)
i.merge(j)
k.merge(l)
for i in range(0, 2):
    for cell in table4.columns[i].cells:
        cell.width = Cm(2.9)
for i in range(2, 14):
    for cell in table4.columns[i].cells:
        cell.width = Cm(1.0)
table4.style = 'MyTable'
docx.add_paragraph()
docx.add_paragraph('由于PCC点的短路容量不同于假定基准最小短路容量,应按照国标附录B进行换算,换算公式如下：')
docx.add_picture('./detect_pic/MVA.png', height=Cm(1.27), width=Cm(2.14))
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
docx.add_paragraph('式中:Sk1：公共连接点的最小短路容量,MVA；\n\t\tSk2：基准短路容量,MVA；\n\t\tIhp：表2中的第h次'
                   '谐波电流允许值,A；\n\t\tIh：短路容量为Sk1时的第h次谐波电流允许值,A。')
docx.add_paragraph('按国标附录C的要求,在公共连接点处第i个用户的第h次谐波电流允许值按下式进行换算：')
docx.add_picture('./detect_pic/Ihi.png', height=Cm(0.74), width=Cm(3.28))
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
docx.add_paragraph('式中:Ih：附录B换算的第h次谐波电流允许值,A；\n\t\tSi：第i个用户的用电协议容量,MVA；\n\t\tSt：公共连接'
                   '点的供电设备容量,MVA；\n\t\tα：相位叠加系数，按表3取值。')
list3 = docx.add_paragraph('\n').add_run('表3: 相位叠加系数')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
list3.font.bold = True
table5 = docx.add_table(rows=1, cols=7)
thead5 = ('谐波次数', '3', '5', '7', '11', '13', '9|>13|偶次')
hdr_cells5 = table5.rows[0].cells
for i in range(0, 7):
    hdr_cells5[i].text = thead5[i]
    hdr_cells5[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells5[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
tr5 = ('α', '1.1', '1.2', '1.4', '1.8', '1.9', '2')
table5_add1 = table5.add_row().cells
for i in range(0, 7):
    table5_add1[i].text = tr5[i]
    table5_add1[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table5_add1[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table5.style = 'MyTable'
docx.add_heading('3. 测试仪器', level=1)
docx.add_paragraph('本次测试采用FLUKE435Ⅱ系列电能质量和能量进行测试。设备参数如下表：')
docx.add_picture('./detect_pic/fluke.png', height=Cm(8.55), width=Cm(8.86))
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
pic1 = docx.add_paragraph().add_run('图1：电能质量分析仪 FLUKE435Ⅱ')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
pic1.font.bold = True
list4 = docx.add_paragraph().add_run('表4: 设备参数表')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
list4.font.bold = True
records2 = (('电压输入', ''),
            ('输入数量', '4 个（3 相 + 中性线）直流耦合'),
            ('最大输入电压', '1000 Vrms'),
            ('标称电压范围', '根据 IEC61000-4-30 可选择 1 V 至 1000 V'),
            ('最大峰值测量电压', '6 kV（仅瞬态模式）'),
            ('输入阻抗', '4 MΩ // 5 pF'),
            ('带宽', '大于 10 kHz，瞬态显示时最高可达 100 kHz'),
            ('缩放比例', '1:1、10:1、100:1、1000:1、10,000:1 及变量'),
            ('电流输入', ''),
            ('输入数量', '个（3 相 + 中性线）直流或交流耦合'),
            ('类型', '夹式变流器，带 mV/A 输出或 i430flex-TF'),
            ('标称输入范围', '0 - ± 3.0 Vpeak，0 - 3.97 Vrms 正弦波（选择 x1，交流+直流耦合）\n\t'
                       '0 - ± 0.3 Vpeak，0 - 0.397 Vrms 正弦波（选择 x10，交流耦合）'),
            ('温度范围', '0.5 Arms - 600 Arms，包括 i430flex-TF（灵敏度 10x）\n\t'
                        '5 Arms - 6000 Arms，包括 i430flex-TF（灵敏度 1x）\n\t'
                        '0.1 mV/A - 1 V/A，以及定制可选的交流或直流钳夹\n\t'
                        '注意：敏感度 x10 意味着分辨率更高，但同时范围缩小。只支持交流电\n\t信号；直流分量被阻隔。'),
            ('阻抗', '1 MΩ')
           )
table6_1 = docx.add_table(rows=0, cols=2)
table6_1.style = 'Table Grid'
for state, content in records2:
    row_cells = table6_1.add_row().cells
    row_cells[0].paragraphs[0].add_run(state).font.size = Pt(9)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.63) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run(content).font.size = Pt(9)
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
for cell0, cell1 in zip(table6_1.columns[0].cells, table6_1.columns[1].cells):
        cell0.width = Cm(3.94)
        cell1.width = Cm(11.09)
docx.add_page_break()
table6_2 = docx.add_table(rows=0, cols=4)
records3_1 = (('带宽', '> 10 kHz'), ('缩放比例', '1:1、10:1、100:1、1000:1、10,000:1 及变量'),
            ('标称频率', '50 Hz，60 Hz'), ('分辨率', '8 信道 16 位模数转换器'),
            ('最高取样速度', '每个信道同时取样时为 200 kS/s'),
            ('RMS 取样', '10/121 个周期为 5000 个样本（依照 IEC 61000-4-30）'),
            ('PLL 同步', '10/121 个周期为 4096 个样本（依照 IEC 61000-4-7）'))
records3_2 = ('电压/电流/频率', '测量范围', '分辨率', '精度')
records3_3 = (('电压 %f', '0.0 … 100.0%', '0.1%', '± 0.1% ± n x 0.1%'),
            ('电压 %r', '0.0 … 100.0%', '0.1%', '± 0.1% ± n x 0.4%'),
            ('绝对电压', '0.0 … 1000 V', '0.1 V', '± 5% (*)'), ('电压总谐波失真 (THD)', '0.0 … 100.0%', '0.1%', '± 2.5%'),
            ('电流 %f', '0.0 … 100.0%', '0.1%', '± 0.1% ± n x 0.1%'),
            ('绝对电流', '0.0 … 600 A', '0.1 A', '± 5% ± 5 个计数'),
            ('电流总谐波失真 (THD)', '0.0 … 100.0%', '0.1%', '± 2.5%'),
            ('功率 %f 或 %r', '0.0 … 100.0%', '0.1%', '± n x 2%'),
            ('绝对功率', '取决于钳夹缩放和标称电压', '', '± 5% ± n x 2% ± 10 个\n\t计数'),
            ('功率总谐波失真 (THD)', '0.0 … 100.0%', '0.1%', '± 5%'), ('相位', '-360º ... +0º', '1º', '± n × 1º (8)'),
            ('闪变Plt，Pst，Pst（1\n\t 分钟）Pinst', '0.00 … 20.00', '0.01', '± 5%'))
records3_4 = (
            ('标准', ''),
            ('所用测量方法', 'IEC61000-4-30 第 2 版 A 级'),
            ('测量性能', 'Fluke 435-II/437-II IEC61000-4-30 A 级'), ('电能质量', 'EN50160'),
            ('闪变', 'IEC 61000-4-15'), ('谐波', 'IEC 61000-4-7'),
            ('Shipboard V/A/Hz', 'MIL-STD-1399-300B'))

table6_2.style = 'Table Grid'

for state, content1 in records3_1:
    row_cells = table6_2.add_row().cells
    row_cells[0].paragraphs[0].add_run(state).font.size = Pt(9)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.63) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run(content1).font.size = Pt(9)
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
addr2 = table6_2.add_row().cells
addr2[0].paragraphs[0].add_run('谐波次数 (n)').font.size = Pt(9)
addr2[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
addr2[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)

addr2_5 = table6_2.add_row().cells
addr2_5[0].paragraphs[0].add_run(records3_2[0]).font.size = Pt(9)
addr2_5[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
addr2_5[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.63) # 小五字体长度0.317cm
addr2_5[1].paragraphs[0].add_run(records3_2[1]).font.size = Pt(9)
addr2_5[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
addr2_5[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
addr2_5[2].paragraphs[0].add_run(records3_2[2]).font.size = Pt(9)
addr2_5[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
addr2_5[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
addr2_5[3].paragraphs[0].add_run(records3_2[3]).font.size = Pt(9)
addr2_5[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
addr2_5[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
addr3 = table6_2.add_row().cells
addr3[0].paragraphs[0].add_run('谐波次数 (n)').font.size = Pt(9)
addr3[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
addr3[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
addr3[2].paragraphs[0].add_run('直流，1..50 分组：谐波依照 IEC61000-4-7 分组').font.size = Pt(9)
addr3[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
addr3[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
addr4 = table6_2.add_row().cells
addr4[0].paragraphs[0].add_run('谐间波次数').font.size = Pt(9)
addr4[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
addr4[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
addr4[2].paragraphs[0].add_run('关闭 (Off)，1..50 分组：谐波和谐间波依照\n\tIEC61000-4-7 分组').font.size = Pt(9)
addr4[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
addr4[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
for state, content1, content2, content3 in records3_3:
    row_cells = table6_2.add_row().cells
    row_cells[0].paragraphs[0].add_run(state).font.size = Pt(9)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.63) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run(content1).font.size = Pt(9)
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
    row_cells[2].paragraphs[0].add_run(content2).font.size = Pt(9)
    row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
    row_cells[3].paragraphs[0].add_run(content3).font.size = Pt(9)
    row_cells[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
for state, content1 in records3_4:
    row_cells = table6_2.add_row().cells
    row_cells[0].paragraphs[0].add_run(state).font.size = Pt(9)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.63) # 小五字体长度0.317cm
    row_cells[2].paragraphs[0].add_run(content1).font.size = Pt(9)
    row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
addr5 = table6_2.add_row().cells
addr5[0].paragraphs[0].add_run('安全性').font.size = Pt(14)
addr5[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
addr5[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.99)

table6_2.cell(0, 1).merge(table6_2.cell(0, 3))
table6_2.cell(1, 1).merge(table6_2.cell(1, 3))
table6_2.cell(2, 1).merge(table6_2.cell(2, 3))
table6_2.cell(3, 1).merge(table6_2.cell(3, 3))
table6_2.cell(4, 1).merge(table6_2.cell(4, 3))
table6_2.cell(5, 1).merge(table6_2.cell(5, 3))
table6_2.cell(6, 1).merge(table6_2.cell(6, 3))
table6_2.cell(7, 0).merge(table6_2.cell(7, 3))
table6_2.cell(9, 0).merge(table6_2.cell(9, 1))
table6_2.cell(9, 2).merge(table6_2.cell(9, 3))
table6_2.cell(10, 0).merge(table6_2.cell(10, 1))
table6_2.cell(10, 2).merge(table6_2.cell(10, 3))
table6_2.cell(19, 1).merge(table6_2.cell(19, 2))
table6_2.cell(23, 1).merge(table6_2.cell(23, 3))
table6_2.cell(24, 0).merge(table6_2.cell(24, 1))
table6_2.cell(24, 2).merge(table6_2.cell(24, 3))
table6_2.cell(25, 0).merge(table6_2.cell(25, 1))
table6_2.cell(25, 2).merge(table6_2.cell(25, 3))
table6_2.cell(26, 0).merge(table6_2.cell(26, 1))
table6_2.cell(26, 2).merge(table6_2.cell(26, 3))
table6_2.cell(27, 0).merge(table6_2.cell(27, 1))
table6_2.cell(27, 2).merge(table6_2.cell(27, 3))
table6_2.cell(28, 0).merge(table6_2.cell(28, 1))
table6_2.cell(28, 2).merge(table6_2.cell(28, 3))
table6_2.cell(29, 0).merge(table6_2.cell(29, 1))
table6_2.cell(29, 2).merge(table6_2.cell(29, 3))
table6_2.cell(30, 0).merge(table6_2.cell(30, 3))
for i in range(0, 4):
    for cell in table6_2.columns[i].cells:
        cell.width = Cm(3.75)
docx.add_page_break()
table6_3 = docx.add_table(rows=0, cols=2)
records3_5 = (('符合标准', 'IEC/EN61010-1-2001,\n\tCAN/CSA C22.2 第 61010-1-04 号（包括\n\tcCSAus 认证），\n\t'
                          'UL 第 61010-1 号，\n\t测量控制和实验室用电气设备的安全要求，第 1 \n\t部分：一般要求。\n\t'
                          '额定：600 V 第四类 (CAT IV) 1000 V 第三类\n\t(CAT III) 污染等级 2'),
              ('橡胶输入端最大电压', '1000 V 第三类 (CAT III) / 600 V 第四类 (CAT IV)'),
              ('电流 BNC 输入端的最大电压', '30 V 最大值'))
table6_3.style = 'Table Grid'
for state, content1 in records3_5:
    row_cells = table6_3.add_row().cells
    row_cells[0].paragraphs[0].add_run(state).font.size = Pt(9)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.63) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run(content1).font.size = Pt(9)
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
# for s in docx.paragraphs[7:16]:
#     print(s.text) # 检查段落
docx.add_heading('4 测试参数', level=1)
docx.add_paragraph('本次测试主要以谐波测为主，包括2～50次谐波。同时兼顾电压偏差、功率因数等其他电能质量指标。')
docx.add_heading('5 测试参数', level=1)
docx.add_heading('5.1 测试位置', level=2)
docx.add_picture('./detect_pic/test_position.png', height=Cm(7.04), width=Cm(14.63))
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
pic2 = docx.add_paragraph().add_run('图2：测试位置')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
pic2.font.bold = True
docx.add_heading('5.2 测试接线方式：三相四线', level=2)
docx.add_picture('./detect_pic/spot_wiring.png', height=Cm(4.63), width=Cm(7.94))
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
pic3 = docx.add_paragraph().add_run('图3：现场接线示意图')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
pic3.font.bold = True
docx.add_heading('6 测试结果及其分析', level=1)
docx.add_heading('6.1 电压', level=2)
docx.add_heading('6.1.1 电压有效值变化趋势图', level=3)
docx.add_paragraph('测试期间，________________，三相电压变化趋势图如下图所示（线电压）')
docx.add_picture('./detect_pic/Virms_line.png', height=Cm(8.28), width=Cm(14.63))
pic4 = docx.add_paragraph().add_run('图4：电压有效值变化趋势图')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
pic4.font.size = Pt(10.5)
pic4.font.bold = True
docx.add_heading('6.1.2 电压有效值', level=3)
list5 = docx.add_paragraph().add_run('表5：电压幅值测试结果（单位：V）')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
list5.font.bold = True
records4 = (('项目名称', '最大值', '最小值', '95%概率值'),
            ('A相电压有效值', df['Vrms ph-ph AB Avg'].max(), df['Vrms ph-ph AB Avg'].min(),
             np.percentile(df['Vrms ph-ph AB Avg'], 95, axis=0)),
            ('B相电压有效值', df['Vrms ph-ph BC Avg'].max(), df['Vrms ph-ph BC Avg'].min(),
             np.percentile(df['Vrms ph-ph BC Avg'], 95, axis=0)),
            ('C相电压有效值', df['Vrms ph-ph CA Avg'].max(), df['Vrms ph-ph CA Avg'].min(),
             np.percentile(df['Vrms ph-ph CA Avg'], 95, axis=0)))
table7 = docx.add_table(rows=0, cols=4)
for nam, maxv, minv, perc in records4:
    row_cells = table7.add_row().cells
    row_cells[0].paragraphs[0].add_run(nam)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.74) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run(str(maxv))
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[2].paragraphs[0].add_run(str(minv))
    row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[3].paragraphs[0].add_run(str(perc))
    row_cells[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table7.style = 'MyTable'
docx.add_heading('6.2 电压偏差', level=2)
docx.add_heading('6.2.1电压偏差值', level=3)
list6 = docx.add_paragraph().add_run('表6：电压偏差测试结果（单位：％）')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
list6.font.bold = True
th5 = ('项目名称', '最大值', '最小值', '95%概率值')
records5 = (('A相电压正偏差', (df['Vrms ph-ph AB Avg'].max() / 380 - 1) * 100,
             abs(df['Vrms ph-ph AB Avg'].min() / 380 - 1) * 100,
             abs(np.percentile(df['Vrms ph-ph AB Avg'], 95, axis=0) / 380 - 1) * 100),
            ('B相电压正偏差', (df['Vrms ph-ph BC Avg'].max() / 380 - 1) * 100,
                         abs(df['Vrms ph-ph BC Avg'].min() / 380 - 1) * 100,
             abs(np.percentile(df['Vrms ph-ph BC Avg'], 95, axis=0) / 380 - 1) * 100),
            ('C相电压正偏差', (df['Vrms ph-ph CA Avg'].max() / 380 - 1) * 100,
            abs(df['Vrms ph-ph CA Avg'].min() / 380 - 1) * 100,
             abs(np.percentile(df['Vrms ph-ph CA Avg'], 95, axis=0)/ 380 - 1) * 100))
table8 = docx.add_table(rows=1, cols=4)
hdr_cells8 = table8.rows[0].cells
for i in range(0, 4):
    hdr_cells8[i].text = th5[i]
    hdr_cells8[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells8[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
for nam, maxv, minv, perc in records5:
    row_cells = table8.add_row().cells
    row_cells[0].paragraphs[0].add_run(nam)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.74) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run('%.2f' % float(maxv))
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[2].paragraphs[0].add_run('%.2f' % float(minv))
    row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[3].paragraphs[0].add_run('%.2f' % float(perc))
    row_cells[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table8.style = 'MyTable'
docx.add_heading('6.2.2分析结论', level=3)
docx.add_paragraph('根据国家标准，380V电压值偏差范围应该在±7％范围内，分析上述数据，'
                   '可得出下述结论：A、B、C三相最大电压_____超出国家标准7％范围，电压______。')
docx.add_page_break()
docx.add_heading('6.3 电压谐波', level=2)
docx.add_heading('6.3.1电容______电压谐波总含有率THDU变化趋势', level=3)
docx.add_paragraph('测试期间，无功补偿电容器______，THDU变化趋势图如下图所示')
docx.add_picture('./detect_pic/THDU.png', width=Cm(14.63), height=Cm(7.91))
pic5 = docx.add_paragraph().add_run('图5：电压谐波THDU总含有率变化趋势图')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
pic5.font.size = Pt(10.5)
pic5.font.bold = True
docx.add_heading('6.3.2电容器______电压谐波总含有率值', level=3)
list7 = docx.add_paragraph().add_run('表7：THDU测试结果（单位：%）')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
list7.font.bold = True
th11 = ('项目名称', '最大值', '最小值', '平均值', '95%概率值')
table9 = docx.add_table(rows=1, cols=5)
hdr_cells9 = table9.rows[0].cells
records6 = (
            ('A相电压谐波总含有率', df['THD V AN Avg'].max(), df['THD V AN Avg'].min(), df['THD V AN Avg'].mean(),
             np.percentile(df['THD V AN Avg'], 95, axis=0)),
            ('B相电压谐波总含有率', df['THD V BN Avg'].max(), df['THD V BN Avg'].min(), df['THD V BN Avg'].mean(),
             np.percentile(df['THD V BN Avg'], 95, axis=0)),
            ('C相电压谐波总含有率', df['THD V CN Avg'].max(), df['THD V CN Avg'].min(), df['THD V CN Avg'].mean(),
             np.percentile(df['THD V CN Avg'], 95, axis=0)))
for i in range(0, 5):
    hdr_cells9[i].text = th11[i]
    hdr_cells9[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells9[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
for nam, maxv, minv, meanv, perc in records6:
    row_cells = table9.add_row().cells
    row_cells[0].paragraphs[0].add_run(nam)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.74) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run('%.2f' % float(maxv))
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[2].paragraphs[0].add_run('%.2f' % float(minv))
    row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[3].paragraphs[0].add_run('%.2f' % float(meanv))
    row_cells[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[4].paragraphs[0].add_run('%.2f' % float(perc))
    row_cells[4].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[4].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table9.style = 'MyTable'
wid = (4.7, 2.45, 2.45, 2.45, 2.45)
for i in range(0, 5):
    for cell in table9.columns[i].cells:
        cell.width = Cm(wid[i])
docx.add_heading('6.3.2 主要分次电压谐波含有率变化趋势图', level=3)
Ulist = []
sum = 0
for files in os.walk('./detect_pic/'):
    for i in range(2, 13):
        if 'Vharm' + str(i) + '.png' in files[2]:
            sum += 1
            Ulist.append(i)
            docx.add_picture('./detect_pic/' + 'Vharm' + str(i) + '.png', width=Cm(14.63), height=Cm(8.07))
            pic_num = docx.add_paragraph().add_run('图%d： %d次电压谐波含有率' % ((5 + sum), i))
            docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pic_num.font.size = Pt(10.5)
            pic_num.font.bold = True
            if sum % 2 == 0:
                docx.add_page_break()

list8 = docx.add_paragraph().add_run('表8：测试结果（%）')
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
list8.font.bold = True
th10 = ('项目名称', '最大值', '最小值', '平均值')

table10 = docx.add_table(rows=1, cols=4)
for i in range(0, 4):
    table10.rows[0].cells[i].text = th10[i]
    table10.rows[0].cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table10.rows[0].cells[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
for i in Ulist:
    cols_a = 'Volts Harmonics' + str(i) + ' AN Avg'
    cols_b = 'Volts Harmonics' + str(i) + ' BN Avg'
    cols_c = 'Volts Harmonics' + str(i) + ' CN Avg'
    maxValue = np.max([df[cols_a], df[cols_b], df[cols_c]])
    minValue = np.min([df[cols_a], df[cols_b], df[cols_c]])
    meanValue = np.mean([df[cols_a], df[cols_b], df[cols_c]])
    row_cells = table10.add_row().cells
    row_cells[0].paragraphs[0].add_run(str(i) + '次电压谐波含有率')
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.74) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run('%.2f' % float(maxValue))
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[2].paragraphs[0].add_run('%.2f' % float(minValue))
    row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[3].paragraphs[0].add_run('%.2f' % float(meanValue))
    row_cells[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)

table10.style = 'MyTable'
docx.add_heading('6.3.4分析结论', level=3)
docx.add_paragraph('根据国标GB/T14549-93的要求，0.38kV级电网公共连接点电压总谐波畸变率限值为5%，'
                   '奇次谐波含有率4％，偶次谐波含有率2％，上述测试数据表明：')
docx.paragraphs[-1].add_run('________________________。').font.bold = True
docx.add_page_break()
docx.add_heading('6.4 电流有效值', level=2)
docx.add_heading('6.4.1电流有效值变化趋势', level=3)
docx.add_paragraph('测试期间，无功补偿电容器投入时，三相电流变化趋势图如下图所示：')
docx.add_picture('./detect_pic/current.png', height=Cm(8.4), width=Cm(14.63))
pic_num = docx.add_paragraph().add_run('图%d：电流有效值变化趋势图' % (sum + 6))
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
pic_num.font.size = Pt(10.5)
pic_num.font.bold = True
docx.add_heading('6.4.2电流有效值', level=3)
list9 = docx.add_paragraph().add_run('表9：电流有效值测试结果（单位：A）').bold = True
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
records7 = (('A相电流有效值', df['Current A Avg'].max(), df['Current A Avg'].min(), df['Current A Avg'].mean(),
            np.percentile(df['Current A Avg'], 95, axis=0)),
            ('B相电流有效值', df['Current B Avg'].max(), df['Current B Avg'].min(), df['Current B Avg'].mean(),
            np.percentile(df['Current B Avg'], 95, axis=0)),
            ('C相电有效值', df['Current C Avg'].max(), df['Current C Avg'].min(), df['Current C Avg'].mean(),
            np.percentile(df['Current C Avg'], 95, axis=0)))

table11 = docx.add_table(rows=1, cols=5)
for i in range(0, 5):
    table11.rows[0].cells[i].text = th11[i]
    table11.rows[0].cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table11.rows[0].cells[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
for nam, maxv, minv, meanv,perc in records7:
    row_cells = table11.add_row().cells
    row_cells[0].paragraphs[0].add_run(nam)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.74) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run('%.2f' % float(maxv))
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[2].paragraphs[0].add_run('%.2f' % float(minv))
    row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[3].paragraphs[0].add_run('%.2f' % float(meanv))
    row_cells[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[4].paragraphs[0].add_run('%.2f' % float(perc))
    row_cells[4].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[4].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
for i in range(0, 5):
    for cell in table11.columns[i].cells:
        cell.width = Cm(wid[i])
table11.style = 'MyTable'
docx.add_page_break()
docx.add_heading('6.4.3 分析结论', level=3)
docx.add_paragraph().add_run('根据变压器容量————kVA计算可得变压器二次侧额定电流———A，测量期'
                             '间变压器95%概率值负荷率—————%，最高负荷—————%，变压器________。')
docx.add_heading('6.5 电流谐波',level=2)
docx.add_heading('6.5.1电流谐波总含有率THDI变化趋势', level=3)
docx.add_paragraph().add_run('测试期间，无功补偿电容器________时，电流谐波总含有率变化趋势图如下图所示')
docx.add_picture('./detect_pic/THDI.png', height=Cm(8.4), width=Cm(14.63))
print(sum)
docx.add_paragraph().add_run('图%d：电流谐波总含有率变化趋势图' % (sum + 7)).bold = True
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
docx.add_heading('6.5.2 电容器______时电流谐波总含有率值', level=3)
list10 = docx.add_paragraph().add_run('表10：电流谐波总含有率（%）').bold = True
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

records8 = (('A相电流谐波总含有率', df['THD A A Avg'].max(), df['THD A A Avg'].min(), df['THD A A Avg'].mean(),
             np.percentile(df['THD A A Avg'], 95, axis=0)),
            ('B相电流谐波总含有率', df['THD A B Avg'].max(), df['THD A B Avg'].min(), df['THD A B Avg'].mean(),
             np.percentile(df['THD A B Avg'], 95, axis=0)),
            ('C相电流谐波总含有率', df['THD A C Avg'].max(), df['THD A C Avg'].min(), df['THD A C Avg'].mean(),
             np.percentile(df['THD A C Avg'], 95, axis=0)))
table12 = docx.add_table(rows=1, cols=5)
for i in range(0, 5):
    table12.rows[0].cells[i].text = th11[i]
    table12.rows[0].cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table12.rows[0].cells[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
for nam, maxv, minv, meanv,perc in records8:
    row_cells = table12.add_row().cells
    row_cells[0].paragraphs[0].add_run(nam).font.size = Pt(9)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.63) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run('%.2f' % float(maxv)).font.size = Pt(9)
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
    row_cells[2].paragraphs[0].add_run('%.2f' % float(minv)).font.size = Pt(9)
    row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
    row_cells[3].paragraphs[0].add_run('%.2f' % float(meanv)).font.size = Pt(9)
    row_cells[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
    row_cells[4].paragraphs[0].add_run('%.2f' % float(perc)).font.size = Pt(9)
    row_cells[4].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[4].paragraphs[0].paragraph_format.left_indent = -Cm(0.63)
for i in range(0, 5):
    for cell in table12.columns[i].cells:
        cell.width = Cm(wid[i])
table12.style = 'MyTable'
ilist = []
docx.add_heading('6.5.3主要分次电流谐波值变化趋势', level=3)
for files in os.walk('./detect_pic/'):
    for i in range(2, 13):
        if 'Iharm' + str(i) + '.png' in files[2]:
            sum += 1
            ilist.append(i)
            docx.add_picture('./detect_pic/' + 'Iharm' + str(i) + '.png', width=Cm(14.63), height=Cm(8.07))
            pic_num = docx.add_paragraph().add_run('图%d： %d次电流谐波值' % ((7 + sum), i))
            docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pic_num.font.size = Pt(10.5)
            pic_num.font.bold = True
            if sum % 2 != 0:
                docx.add_page_break()
list11 = list10 = docx.add_paragraph().add_run('表11：测试结果（A）').bold = True
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
table13 = docx.add_table(rows=1, cols=4)
for i in range(0, 4):
    table13.rows[0].cells[i].text = th10[i]
    table13.rows[0].cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table13.rows[0].cells[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
for i in ilist:
    coln_a = 'Amps Harmonics' + str(i) + ' A Avg'
    coln_b = 'Amps Harmonics' + str(i) + ' B Avg'
    coln_c = 'Amps Harmonics' + str(i) + ' C Avg'
    trans_arrray0 = np.array(df[coln_a] * df['Current A Avg'])
    trans_arrray1 = np.array(df[coln_b] * df['Current B Avg'])
    trans_arrray2 = np.array(df[coln_c] * df['Current C Avg'])

    maxValue = np.max([trans_arrray0, trans_arrray1, trans_arrray2]) / 100
    minValue = np.min([trans_arrray0, trans_arrray1, trans_arrray2]) / 100
    meanValue = np.mean([trans_arrray0, trans_arrray1, trans_arrray2]) / 100
    row_cells = table13.add_row().cells
    row_cells[0].paragraphs[0].add_run(str(i) + '次电流谐波值')
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.74) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run('%.2f' % float(maxValue))
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[2].paragraphs[0].add_run('%.2f' % float(minValue))
    row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[3].paragraphs[0].add_run('%.2f' % float(meanValue))
    row_cells[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table13.style = 'MyTable'
docx.add_heading('6.5.4分析结论', level=3)
docx.add_paragraph('依据GB/T14549-93电能质量 公用电网谐波国家标准，公共连接点的全部用户向该点注入的谐波电流允许值如下：')
table14 = docx.add_table(rows=1, cols=14)
hdr_cells14 = table14.rows[0].cells

for i in range(0, 3):
    hdr_cells14[i].text = thead3[i]
    hdr_cells14[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells14[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)

table14_add1 = table14.add_row().cells
for i in range(2, 14):
    table14_add1[i].text = tr3_1[i]
    table14_add1[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table14_add1[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table14_add2 = table14.add_row().cells
for i in range(0, 14):
    table14_add2[i].text = tr3_2[i]
    table14_add2[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table14_add2[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
g = table14.cell(0, 2)
h = table14.cell(0, 13)
i = table14.cell(0, 0)
j = table14.cell(1, 0)
k = table14.cell(0, 1)
l = table14.cell(1, 1)
g.merge(h)
i.merge(j)
k.merge(l)
for i in range(0, 2):
    for cell in table14.columns[i].cells:
        cell.width = Cm(2.9)
for i in range(2, 14):
    for cell in table14.columns[i].cells:
        cell.width = Cm(1.0)
table14.style = 'MyTable'
docx.add_paragraph()

table15 = docx.add_table(rows=1, cols=14)
hdr_cells15 = table15.rows[0].cells

for i in range(0, 3):
    hdr_cells15[i].text = thead4[i]
    hdr_cells15[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells15[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)

table15_add1 = table15.add_row().cells
for i in range(2, 14):
    table15_add1[i].text = tr4_1[i]
    table15_add1[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table15_add1[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table15_add2 = table15.add_row().cells
for i in range(0, 14):
    table15_add2[i].text = tr4_2[i]
    table15_add2[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table15_add2[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
g = table15.cell(0, 2)
h = table15.cell(0, 13)
i = table15.cell(0, 0)
j = table15.cell(1, 0)
k = table15.cell(0, 1)
l = table15.cell(1, 1)
g.merge(h)
i.merge(j)
k.merge(l)
for i in range(0, 2):
    for cell in table15.columns[i].cells:
        cell.width = Cm(2.9)
for i in range(2, 14):
    for cell in table15.columns[i].cells:
        cell.width = Cm(1.0)
table15.style = 'MyTable'
docx.add_paragraph('由上表可以得出：')
docx.paragraphs[-1].add_run('___________。').bold = True
docx.add_page_break()
docx.add_heading('6.6__________时功率因数和无功功率', level=2)
docx.add_heading('6.6.1功率因数值变化趋势', level=3)
docx.add_paragraph('测试期间，____________时，三相功率因数变化趋势图如下图所示')
docx.add_picture('./detect_pic/cos_phi.png', height=Cm(7.93), width=Cm(14.63))
docx.add_paragraph().add_run('图%d：功率因素值变化趋势' % (sum + 8)).font.bold = True
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
docx.add_heading('6.6.2功率因数值', level=3)
list12 = docx.add_paragraph().add_run('表12：功率因数值测试结果').font.bold = True
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
table16 = docx.add_table(cols=5, rows=1)
for i in range(0, 5):
    table16.rows[0].cells[i].text = th11[i]
    table16.rows[0].cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table16.rows[0].cells[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
records9 = (('A相功率因数', df['Cos Phi AN Avg'].max(), df['Cos Phi AN Avg'].min(), df['Cos Phi AN Avg'].mean(),
            np.percentile(df['Cos Phi AN Avg'], 95, axis=0)),
            ('B相功率因数', df['Cos Phi BN Avg'].max(), df['Cos Phi BN Avg'].min(), df['Cos Phi BN Avg'].mean(),
            np.percentile(df['Cos Phi BN Avg'], 95, axis=0)),
            ('C相功率因数', df['Cos Phi CN Avg'].max(), df['Cos Phi CN Avg'].min(), df['Cos Phi CN Avg'].mean(),
            np.percentile(df['Cos Phi CN Avg'], 95, axis=0)),
            ('合计', df['Cos Phi Total Avg'].max(), df['Cos Phi Total Avg'].min(), df['Cos Phi Total Avg'].mean(),
            np.percentile(df['Cos Phi Total Avg'], 95, axis=0)))
for nam, maxv, minv, meanv,perc in records9:
    row_cells = table16.add_row().cells
    row_cells[0].paragraphs[0].add_run(nam)
    row_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.74) # 小五字体长度0.317cm
    row_cells[1].paragraphs[0].add_run('%.2f' % float(maxv))
    row_cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[2].paragraphs[0].add_run('%.2f' % float(minv))
    row_cells[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[3].paragraphs[0].add_run('%.2f' % float(meanv))
    row_cells[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
    row_cells[4].paragraphs[0].add_run('%.2f' % float(perc))
    row_cells[4].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    row_cells[4].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table16.style = 'MyTable'
docx.add_page_break()
docx.add_heading('6.6.3无功功率值变化趋势', level=3)
docx.add_paragraph('测试期间，____________时，总无功功率变化趋势图如下图所示')
docx.add_picture('./detect_pic/reactive.png', height=Cm(8.03), width=Cm(14.63))
docx.add_paragraph().add_run('图%d：无功功率变化趋势' % (sum + 9)).font.bold = True
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
docx.add_heading('6.6.4无功功率值', level=3)
list13 = docx.add_paragraph().add_run('表13：无功功率值测试结果（单位：kvar）').font.bold = True
docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
table17 = docx.add_table(rows=1, cols=4)
for i in range(0, 4):
    table17.rows[0].cells[i].text = th10[i]
    table17.rows[0].cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table17.rows[0].cells[i].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
records10 = ('无功功率', df['Reactive Power Total Avg'].max(),
             df['Reactive Power Total Avg'].min(),
             df['Reactive Power Total Avg'].mean())

row_cells2 = table17.add_row().cells
row_cells2[0].paragraphs[0].add_run(records10[0])
row_cells2[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
row_cells2[0].paragraphs[0].paragraph_format.left_indent = -Cm(0.74) # 小五字体长度0.317cm
row_cells2[1].paragraphs[0].add_run('%.2f' % float(records10[1] / 1000))
row_cells2[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
row_cells2[1].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
row_cells2[2].paragraphs[0].add_run('%.2f' % float(records10[2] / 1000))
row_cells2[2].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
row_cells2[2].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
row_cells2[3].paragraphs[0].add_run('%.2f' % float(records10[3] / 1000))
row_cells2[3].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
row_cells2[3].paragraphs[0].paragraph_format.left_indent = -Cm(0.74)
table17.style = 'MyTable'
docx.add_heading('6.6.5分析结论', level=3)
docx.add_paragraph().add_run('________________')
docx.add_page_break()
docx.add_heading('7 总结', level=1)
docx.add_paragraph('通过本次测试结合现场调查： ')
for i in range(0, 5):
    docx.add_paragraph()
docx.add_paragraph('以上结论仅针对测试时段内数据分析结果及现场用电情况')
docx.add_paragraph('本次测试的数据已超出国家对电流谐波含量的规定，已经属于谐波污染场所。应当加装必要的滤波设备'
                   '，控制电网谐波含有率。这种超出规定的情况下，对电网中的用电设备会造成如下常见的危害：')
docx.add_paragraph('第一：导致电缆发热；\n\t第二：导致变压器容量降低、过热、异响；\n\t第三：导致变无功补偿装置损坏；\n\t'
                   '第四：加速电子设备的老化，缩短使用寿命；\n\t第五：继电保护装置误动作；\n\t第六：导致额外的电'
                   '能损耗，电费成本增加；\n\t第七：造成采样误差偏大，控制失效等不良情况。')
docx.add_paragraph('并且从无功补偿方面进行考虑会出现：\n\t').add_run('1、电容器会对谐波电流进行放大，对电网产生二次污染。\n\t'
                                                    '2、电容器本身出现电流过载现象，影响使用寿命甚至造成损坏。\n\t'
                                                    '3、电容投切器件也会因为谐波影响，造成触点寿寿损伤及烧结现象，'
                                                    '导致无法正常运行。\n\t4、当前普通无功补偿对谐波抑制能力不足，'
                                                    '会出现采样偏差大，抑制失效等不良情况。\n\t')
docx.add_paragraph('本次测试电网存在三相负荷不平衡，电网中的不平衡电流会增加线路及变压器的铜损、还会增加变压器'
                   '的铁损，降低变压器的输出能力甚至会影响变压器的安全运行，造成如下常见的危害：')
docx.add_paragraph('1.增加线路的电能损耗。\n\t2.增加配电变压器的电能损耗。\n\t3.配变出力减少。\n\t4.配变产生零序电流。\n\t'
                   '5.影响用电设备的安全运行。\n\t6.变压器、铜排异响变色等现象')
docx.add_paragraph('浙江南德电气有限公司技术部\n').add_run(text=now.strftime('%Y') + '年' + now.strftime('%m') + '月' +
                                              now.strftime('%d') + '日')
docx.paragraphs[-1].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
docx.save('./text/'+ COMPANY + '.docx')

